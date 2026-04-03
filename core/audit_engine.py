import docx
import re

def extract_legal_references(input_path):
    """
    Truy quét toàn diện file Word bao gồm: Paragraphs, Tables, Headers, Footers.
    Sử dụng Regex Unicode để không bỏ sót văn bản tiếng Việt.
    """
    doc = docx.Document(input_path)
    full_text = []

    # 1. Quét nội dung chính
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # 2. Quét bảng biểu
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    # 3. Quét Header & Footer (Nơi thường để căn cứ pháp lý)
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            full_text.append(header_para.text)
        for footer_para in section.footer.paragraphs:
            full_text.append(footer_para.text)
    
    text = "\n".join(full_text)
    
    # Bộ Regex "Siêu lọc" - Cực kỳ linh hoạt với mọi cấu trúc tiếng Việt
    patterns = [
        # 1. Nghị định, Thông tư, Quyết định, Nghị quyết, Văn bản hợp nhất
        # Bắt từ khóa -> Số hiệu -> Hậu tố (Hỗ trợ Unicode như NĐ-CP, TT-BXD...)
        r"(?:Nghị định|Thông tư|Quyết định|Nghị quyết|VBHN)(?:\s+số)?\s+\d{1,5}/\d{4}/[^\s,;]+",
        
        # 2. Luật kèm theo tên và số hiệu (Hỗ trợ khoảng cách xa và dấu phẩy ở giữa)
        # Ví dụ: Luật Xây dựng ... số 50/2014/QH13
        r"Luật[\s\wÀ-Ỹà-ỹ,]+?số\s+\d{1,5}/\d{4}/QH\d{1,2}\b",
        
        # 3. Luật trực tiếp (Ví dụ: Luật số 10/2012/QH13)
        r"Luật(?:\s+số)?\s+\d{1,5}/\d{4}/QH\d{1,2}\b",
        
        # 4. Luật theo tên và năm (Ví dụ: Luật Xây dựng 2014)
        r"Luật\s+[A-ZÀ-Ỹ][a-zà-ỹ\s\w]+[12]\d{3}\b",
        
        # 5. Tiêu chuẩn, Quy chuẩn (TCVN, QCVN, TCXDVN, TCVN/XD...)
        r"(?:TCVN|TCXDVN|QCVN|TCVN/XD|QCXDVN)\s+\d+[:\-\s][12]\d{3}(?:/[^\s,;.]+)?\b"
    ]
    
    found_refs = []
    for p in patterns:
        matches = re.finditer(p, text, flags=re.IGNORECASE | re.UNICODE)
        for m in matches:
            ref = m.group(0).strip()
            # Làm sạch kỹ hơn các ký tự dư thừa ở đầu/cuối
            ref = re.sub(r'^[-—\t\s]+', '', ref) # Bỏ dấu gạch đầu dòng
            ref = re.sub(r'[\s.;,]+$', '', ref) # Bỏ dấu chấm phẩy cuối
            found_refs.append(ref)
    
    # Loại bỏ trùng lặp và sắp xếp
    unique_refs = sorted(list(set(found_refs)))
    return unique_refs

def audit_legal_status(refs):
    """
    Phân loại nguồn tra cứu chính xác.
    """
    results = []
    for ref in refs:
        # Tiêu chuẩn/Quy chuẩn -> VSQI
        is_standard = any(std in ref.upper() for std in ["TCVN", "QCVN", "TCXDVN", "TCVN/XD", "QCXDVN", "TIÊU CHUẨN", "QUY CHUẨN"])
        
        target_domain = "tieuchuan.vsqi.gov.vn" if is_standard else "thuvienphapluat.vn"
        
        search_query = f"{ref} {target_domain} tình trạng hiệu lực"
        search_link = f"https://www.google.com/search?q={search_query.replace(' ', '+')}"
        results.append({
            "symbol": ref,
            "link": search_link
        })
    return results
